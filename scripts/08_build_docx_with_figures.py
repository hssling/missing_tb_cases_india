"""
Convert a Markdown manuscript to DOCX and embed key figures at locations where
their labels appear in the text.

Usage:
    python scripts/08_build_docx_with_figures.py \
        --md reports/tb_manuscript_v5.md \
        --docx reports/tb_manuscript_v5_with_figures.docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches


FIGURE_MAP = {
    "Figure 1": {
        "path": Path("output/figures/national_trend.png"),
        "caption": "National TB incidence and notifications with missed-case gap."
    },
    "Figure 2": {
        "path": Path("output/figures/detection_boxplot.png"),
        "caption": "Distribution of state detection coverage (2020–2023)."
    },
    "Figure 3": {
        "path": Path("output/figures/missed_cases_density.png"),
        "caption": "Kernel density of state missed cases across years."
    },
    "Figure 4": {
        "path": Path("output/figures/system_vs_detection.png"),
        "caption": "System strength versus detection coverage, coloured by risk score."
    },
    "Figure 5": {
        "path": Path("output/figures/state_detection_map.png"),
        "caption": "Estimated detection coverage across Indian states."
    },
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build DOCX with embedded figures.")
    parser.add_argument("--md", default="reports/tb_manuscript_v5.md", help="Path to Markdown manuscript.")
    parser.add_argument("--docx", default="reports/tb_manuscript_v5_with_figures.docx", help="Output DOCX path.")
    return parser.parse_args()


def parse_table(table_lines: list[str]) -> tuple[list[str], list[list[str]]] | None:
    rows: list[list[str]] = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if set(stripped.replace("|", "").strip()) <= set("-: "):
            continue
        rows.append([cell.strip() for cell in stripped.strip("|").split("|")])
    if not rows:
        return None
    return rows[0], rows[1:]


def maybe_insert_figure(document: Document, line: str, inserted: set[str]) -> None:
    for label, meta in FIGURE_MAP.items():
        if label in line and label not in inserted and meta["path"].exists():
            document.add_picture(str(meta["path"]), width=Inches(6))
            document.add_paragraph(f"{label}. {meta['caption']}")
            inserted.add(label)


def main() -> None:
    args = parse_args()
    md_path = Path(args.md)
    docx_path = Path(args.docx)
    text = md_path.read_text(encoding="utf-8")

    document = Document()
    lines = text.splitlines()
    table_buffer: list[str] = []
    figure_inserted: set[str] = set()

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        parsed = parse_table(table_buffer)
        table_buffer = []
        if not parsed:
            return
        header, data = parsed
        table = document.add_table(rows=1, cols=len(header))
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(header):
            hdr_cells[idx].text = text
        for row in data:
            row_cells = table.add_row().cells
            for idx in range(len(header)):
                row_cells[idx].text = row[idx] if idx < len(row) else ""

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            continue
        flush_table()

        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            document.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        else:
            document.add_paragraph(line)
        maybe_insert_figure(document, stripped, figure_inserted)

    flush_table()
    document.save(docx_path)
    print(f"[08_build_docx_with_figures] Saved {docx_path}")


if __name__ == "__main__":
    main()
