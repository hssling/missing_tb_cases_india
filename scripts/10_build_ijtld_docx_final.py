"""
Convert a Markdown manuscript to DOCX with tables and figures at the end,
and proper superscript formatting for references.

Usage:
    python scripts/10_build_ijtld_docx_final.py \
        --md reports/tb_manuscript_v8_ijtld_beautiful.md \
        --docx reports/tb_manuscript_v8_ijtld_beautiful_final.docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches


FIGURE_MAP = {
    "Figure 1": {
        "path": Path("output/figures/national_trend.png"),
        "caption": "National TB incidence and notifications with missed-case gap."
    },
    "Figure 2": {
        "path": Path("output/figures/detection_boxplot.png"),
        "caption": "Distribution of state detection coverage (2020–2023)."
    },
    "Figure 3": {
        "path": Path("output/figures/missed_cases_density.png"),
        "caption": "Kernel density of state missed cases across years."
    },
    "Figure 4": {
        "path": Path("output/figures/system_vs_detection.png"),
        "caption": "System strength versus detection coverage, coloured by risk score."
    },
    "Figure 5": {
        "path": Path("output/figures/state_detection_map.png"),
        "caption": "Estimated detection coverage across Indian states."
    },
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build DOCX with tables/figures at end and superscripts.")
    parser.add_argument("--md", default="reports/tb_manuscript_v8_ijtld_beautiful.md", help="Path to Markdown manuscript.")
    parser.add_argument("--docx", default="reports/tb_manuscript_v8_ijtld_beautiful_final.docx", help="Output DOCX path.")
    return parser.parse_args()


def parse_table(table_lines: list[str]) -> tuple[list[str], list[list[str]]] | None:
    rows: list[list[str]] = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if set(stripped.replace("|", "").strip()) <= set("-: "):
            continue
        rows.append([cell.strip() for cell in stripped.strip("|").split("|")])
    if not rows:
        return None
    return rows[0], rows[1:]


def add_paragraph_with_superscripts(document: Document, text: str) -> None:
    """Add a paragraph with superscript handling for ^text^."""
    para = document.add_paragraph()
    # Split by ^ to handle superscripts
    parts = re.split(r'(\^\^?[^*]+\^\^?)', text)
    for part in parts:
        if part.startswith('^') and part.endswith('^') and len(part) > 2:
            # Superscript
            superscript_text = part[1:-1]
            run = para.add_run(superscript_text)
            run.font.superscript = True
        else:
            para.add_run(part)


def main() -> None:
    args = parse_args()
    md_path = Path(args.md)
    docx_path = Path(args.docx)
    text = md_path.read_text(encoding="utf-8")

    document = Document()
    lines = text.splitlines()
    table_buffer: list[str] = []
    figures_to_add: list[str] = []  # Collect figure labels to add at end

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        parsed = parse_table(table_buffer)
        table_buffer = []
        if not parsed:
            return
        header, data = parsed
        table = document.add_table(rows=1, cols=len(header))
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(header):
            hdr_cells[idx].text = text
        for row in data:
            row_cells = table.add_row().cells
            for idx in range(len(header)):
                row_cells[idx].text = row[idx] if idx < len(row) else ""

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            continue
        flush_table()

        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            document.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        else:
            add_paragraph_with_superscripts(document, line)
            # Check for figures, but don't insert yet
            for label in FIGURE_MAP:
                if label in stripped and label not in figures_to_add:
                    figures_to_add.append(label)

    flush_table()

    # Now add figures at the end
    if figures_to_add:
        document.add_heading("Figures", level=2)
        for label in figures_to_add:
            meta = FIGURE_MAP[label]
            if meta["path"].exists():
                document.add_picture(str(meta["path"]), width=Inches(6))
                document.add_paragraph(f"{label}. {meta['caption']}")

    document.save(docx_path)
    print(f"[10_build_ijtld_docx_final] Saved {docx_path}")


if __name__ == "__main__":
    main()