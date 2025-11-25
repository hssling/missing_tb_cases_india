"""
Generate a DOCX version of the validation/explainer document with key equations
rendered in bold monospace style.

Usage:
    python scripts/09_build_validation_doc.py \
        --md reports/research_process_validation_v2.md \
        --docx reports/research_process_validation_v2.docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build validation DOCX with formatted formulas.")
    parser.add_argument("--md", default="reports/research_process_validation_v2.md", help="Path to Markdown file.")
    parser.add_argument("--docx", default="reports/research_process_validation_v2.docx", help="Output DOCX path.")
    return parser.parse_args()


def parse_table(table_lines: list[str]) -> tuple[list[str], list[list[str]]] | None:
    rows: list[list[str]] = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if set(stripped.replace("|", "").strip()) <= set("-: "):
            continue
        rows.append([cell.strip() for cell in stripped.strip("|").split("|")])
    if not rows:
        return None
    return rows[0], rows[1:]


def add_formula(document: Document, text: str) -> None:
    para = document.add_paragraph()
    run = para.add_run(text)
    run.font.bold = True
    run.font.name = "Courier New"
    run.font.size = Pt(11)


def main() -> None:
    args = parse_args()
    md_path = Path(args.md)
    docx_path = Path(args.docx)
    text = md_path.read_text(encoding="utf-8")
    document = Document()
    lines = text.splitlines()
    table_buffer: list[str] = []

    formula_lines = {
        "SystemStrength = 0.30·",
        "SystemStrength_z =",
        "Risk = mean(",
        "Risk_z =",
        "logit(p_{s,t})",
        "Σ_s [ notif_{s,t} / p_{s,t} ]",
        "Î_{s,t} =",
        "Missed cases_{s,t} =",
    }

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        parsed = parse_table(table_buffer)
        table_buffer = []
        if not parsed:
            return
        header, data = parsed
        table = document.add_table(rows=1, cols=len(header))
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(header):
            hdr_cells[idx].text = text
        for row in data:
            row_cells = table.add_row().cells
            for idx in range(len(header)):
                row_cells[idx].text = row[idx] if idx < len(row) else ""

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            document.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        elif any(key in stripped for key in formula_lines):
            add_formula(document, stripped)
        else:
            document.add_paragraph(line)

    flush_table()
    document.save(docx_path)
    print(f"[09_build_validation_doc] Saved {docx_path}")


if __name__ == "__main__":
    main()
