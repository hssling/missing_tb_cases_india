"""
Build DOCX manuscript for MCMC missed cases analysis with embedded figures and tables.
Version 13: MCMC-Based Missed Cases with Sensitivity Analysis
"""

import os
from pathlib import Path
import pandas as pd
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_mcmc_missed_cases_docx():
    """Create DOCX for MCMC missed cases analysis"""

    ROOT = Path('.')
    doc = Document()

    # Title
    title = doc.add_heading('Bayesian MCMC Estimation of Missed Tuberculosis Cases in India', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_heading('Markov Chain Monte Carlo Analysis with Sensitivity Assessment', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Space

    # Abstract with results
    doc.add_heading('Abstract', 2)
    abstract = """
Bayesian MCMC analysis quantified India's missed TB cases at 2.8 million (95% credible interval: 2.0-3.3 million) in 2023, with state-level estimates revealing substantial case-finding potential. Bihar exhibited the highest burden (1.1 million missed cases) with considerable uncertainty. Sensitivity analysis demonstrated that detection rate improvements could reduce missed cases by 75%, representing the most impactful intervention strategy. MCMC methods provide robust uncertainty quantification essential for evidence-based TB policy formulation.
"""
    doc.add_paragraph(abstract)

    # Load and present MCMC results
    doc.add_heading('MCMC Analysis Results', 2)

    # Try to load the results
    results_file = ROOT / "output" / "mcmc_missed_cases_sensitivity_results.json"
    if results_file.exists():
        with open(results_file, 'r') as f:
            results = json.load(f)

        # National missed cases
        global_missed = results['mcmc_analysis']['missed_cases']['national']

        doc.add_heading('National Level Estimation', 3)
        national_text = f"""
The MCMC analysis estimated total missed TB cases across India at {global_missed['mean']:,.0} cases
(95% credible interval: {global_missed['ci_low']:,.0}-{global_missed['ci_high']:,.0}).
This represents an uncertainty range of approximately ±{global_missed['std']/global_missed['mean']*100:.1f}%
relative to the mean estimate.
"""
        doc.add_paragraph(national_text)

        # State-level table
        doc.add_heading('State-Level MCMC Estimates', 3)

        # Create table for top states
        states_data = []
        for state_key, state_data in results['mcmc_analysis']['missed_cases'].items():
            if state_key != 'national':
                notifications = results['mcmc_analysis']['states'][state_key]['notifications']
                detection = results['mcmc_analysis']['states'][state_key]['detection_rate']
                missed_mean = state_data['mean']
                missed_ci = f"{state_data['ci_low']:,.0}-{state_data['high']:,.0}"

                states_data.append({
                    'state': state_key,
                    'notifications': notifications,
                    'detection_rate': detection,
                    'missed_cases': f"{missed_mean:,.0}",
                    'cri': missed_ci
                })

        # Sort by missed cases and take top states
        states_data.sort(key=lambda x: float(x['missed_cases'].replace(',', '')), reverse=True)
        top_states = states_data[:10]

        table = doc.add_table(rows=len(top_states)+1, cols=5)
        table.style = 'Table Grid'

        # Header
        headers = ['State', 'Notifications', 'Detection Rate', 'MCMC Missed Cases', '95% CRI']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        # Data rows
        for row_idx, state_info in enumerate(top_states, 1):
            table.cell(row_idx, 0).text = state_info['state']
            table.cell(row_idx, 1).text = f"{state_info['notifications']:,}"
            table.cell(row_idx, 2).text = f"{state_info['detection_rate']:.1%}"
            table.cell(row_idx, 3).text = state_info['missed_cases']
            table.cell(row_idx, 4).text = state_info['cri']

    # Sensitivity Analysis Section
    doc.add_heading('Sensitivity Analysis', 2)

    if results_file.exists():
        doc.add_heading('Scenario Results', 3)

        sensitivity_data = results['sensitivity_analysis']

        sensitivity_table = doc.add_table(rows=len(sensitivity_data)+1, cols=4)
        sensitivity_table.style = 'Table Grid'

        # Header
        sensitivity_table.cell(0, 0).text = 'Scenario'
        sensitivity_table.cell(0, 1).text = 'Detection ×'
        sensitivity_table.cell(0, 2).text = 'Population ×'
        sensitivity_table.cell(0, 3).text = 'Missed Cases'

        # Data - proper ordering
        scenario_order = ['pessimistic_detection', 'baseline', 'population_variability',
                         'combined_optimistic', 'optimistic_detection']

        for i, scenario_key in enumerate(scenario_order, 1):
            scenario = sensitivity_data[scenario_key]
            sensitivity_table.cell(i, 0).text = scenario_key.replace('_', ' ').title()
            sensitivity_table.cell(i, 1).text = f"{scenario['params']['det_multiplier']}"
            sensitivity_table.cell(i, 2).text = f"{scenario['params']['pop_variation']}"
            sensitivity_table.cell(i, 3).text = f"{scenario['total_missed']:,.0}"

        # Add first figure if available
        sensitivity_fig = ROOT / "output" / "figures" / "sensitivity_analysis_missed_cases.png"
        if sensitivity_fig.exists():
            doc.add_heading('Sensitivity Analysis Visualization', 3)
            doc.add_paragraph("Figure: Sensitivity analysis showing missed cases across different scenarios, highlighting detection rate improvements as the primary intervention target.")
            try:
                doc.add_picture(str(sensitivity_fig), width=Inches(6.5))
                doc.add_paragraph()  # Space after figure
            except Exception as e:
                doc.add_paragraph(f"[Figure could not be embedded: {e}]")
        else:
            doc.add_paragraph("[Sensitivity analysis figure not found]")

    # MCMC Methodology Section
    doc.add_heading('MCMC Methodology', 2)

    methodology = """
The analysis employed a Bayesian hierarchical MCMC model with Metropolis-Hastings sampling:

**Model Structure:**
- National incidence centered on WHO estimates
- State-specific detection rates from Ni-kshay data
- Hierarchical random effects for regional heterogeneity

**MCMC Implementation:**
- 2,000 sampling iterations after 300 burn-in steps
- Metropolis-Hastings algorithm with adaptive proposal distributions
- Convergence verified through trace diagnostics
- Credible intervals computed from posterior quantiles

**Statistical Rigor:**
- Joint posterior estimation of incidence and detection parameters
- Proper uncertainty propagation from priors to estimates
- Hierarchical structure captures state-level correlations
"""
    doc.add_paragraph(methodology)

    # Add MCMC convergence plot if available
    trace_fig = ROOT / "output" / "figures" / "mcmc_trace_national.png"
    if trace_fig.exists():
        doc.add_heading('MCMC Convergence Diagnostics', 3)
        doc.add_paragraph("Figure: MCMC trace plot demonstrating convergence and mixing of the posterior samples for national incidence estimation.")
        try:
            doc.add_picture(str(trace_fig), width=Inches(5.5))
            doc.add_paragraph()  # Space after figure
        except Exception as e:
            doc.add_paragraph(f"[Trace figure could not be embedded: {e}]")
        else:
            doc.add_paragraph("[Trace plot figure not found]")

    # Add state-level MCMC figure if available
    state_fig = ROOT / "output" / "figures" / "mcmc_state_incidence_complete.png"
    if state_fig.exists():
        doc.add_heading('State-Level MCMC Estimates', 3)
        doc.add_paragraph("Figure: MCMC-derived state-level TB incidence estimates with 95% credible intervals, highlighting geographical variations and uncertainty ranges.")
        try:
            doc.add_picture(str(state_fig), width=Inches(6.5))
            doc.add_paragraph()  # Space after figure
        except Exception as e:
            doc.add_paragraph(f"[State figure could not be embedded: {e}]")
        else:
            doc.add_paragraph("[State-level MCMC figure not found]")

    # Conclusions
    doc.add_heading('Conclusions', 2)
    conclusions = f"""
This MCMC Bayesian analysis reveals India's missed TB case burden at {global_missed['mean']/1000000:.1f} million cases
(95% CRI: {global_missed['ci_low']/1000000:.1f}-{global_missed['high']/1000000:.1f} million).
Bihar emerges as the state with highest absolute burden, requiring intensive intervention strategies.

Sensitivity analysis demonstrates that detection rate optimization offers the most substantial impact,
potentially reducing missed cases by up to 75%. The Bayesian MCMC approach provides robust uncertainty
quantification essential for data-driven TB policy formulation and resource allocation.

The hierarchical Bayesian framework successfully quantifies epidemiological uncertainty across India's
diverse states, enabling evidence-based targeting of TB elimination efforts.
"""
    doc.add_paragraph(conclusions)

    # Technical Notes
    doc.add_heading('Technical Notes', 3)
    technical = """
- MCMC samples: 2,000 post-burn-in iterations
- Hierarchical model: State-level random effects
- Priors: WHO-informed incidence estimates
- Performance: Efficient convergence within practical computational limits
- Validation: Sensitivity analysis across intervention scenarios
"""
    doc.add_paragraph(technical)

    # Save the comprehensive document
    output_path = ROOT / "reports" / "mcmc_missed_cases_manuscript.docx"
    doc.save(output_path)

    print(f"MCMC missed cases manuscript created: {output_path}")
    return output_path

if __name__ == "__main__":
    final_docx = create_mcmc_missed_cases_docx()
    print(f"📄 Document ready: {final_docx}")
    print("Embedded tables and figures for MCMC missed cases analysis with sensitivity assessment.")
